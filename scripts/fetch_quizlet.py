#!/usr/bin/env python3
"""
Lấy flashcard từ Quizlet (công khai hoặc riêng tư trong lớp).

Quizlet chặn bot (PerimeterX) và bộ riêng tư yêu cầu đăng nhập + tham gia lớp.
Không thể "bypass" quyền riêng tư — cần tài khoản đã được duyệt vào lớp.

Cách làm (bộ MLN111 riêng tư):
  1. Mở Chrome debug, đăng nhập Quizlet, xin vào lớp "MLN111"
  2. Mở link bộ thẻ — phải thấy nội dung flashcard (không phải trang "riêng tư")
  3. Chạy script với --connect 9222

Cài dependency:
  pip install undetected-chromedriver requests selenium

Ví dụ:
  python scripts/fetch_quizlet.py "https://quizlet.com/vn/1017891468/..." --connect 9222
  python scripts/fetch_quizlet.py 1017891468 -o data/quizlet_mln111.json --connect 9222
  python scripts/fetch_quizlet.py 1017891468 --cookies quizlet_cookies.txt
  python scripts/fetch_quizlet.py "URL" --format docx -o data/mln111_quizlet.docx --connect 9222
  python scripts/fetch_quizlet.py --from-json data/quizlet_mln111.json -o data/mln111_quizlet.docx
"""

from __future__ import annotations

import argparse
import csv
import json
import re
import sys
import time
from http.cookiejar import MozillaCookieJar
from pathlib import Path
from typing import Any
from urllib.parse import unquote

ROOT = Path(__file__).resolve().parent.parent

SET_ID_RE = re.compile(r"(?:quizlet\.com/(?:[\w-]+/)?)?(\d{5,})")
API_BASE = "https://quizlet.com/webapi/3.4/studiable-item-documents"
PER_PAGE = 500

PRIVATE_MARKERS = (
    "chế độ riêng tư",
    "che do rieng tu",
    "private mode",
    "this set is private",
    "học phần này ở chế độ riêng tư",
)
LOGIN_MARKERS = (
    "đăng nhập",
    "log in",
    "sign in",
)


class QuizletAccessError(RuntimeError):
    """Không có quyền xem bộ thẻ (riêng tư / chưa đăng nhập)."""


def page_source_lower(driver) -> str:
    return (driver.page_source or "").lower()


def visible_text_lower(driver) -> str:
    from selenium.webdriver.common.by import By

    try:
        return (driver.find_element(By.TAG_NAME, "body").text or "").lower()
    except Exception:
        return ""


def detect_access_block(driver) -> str | None:
    """Trả về thông báo lỗi nếu trang chặn truy cập (riêng tư, chưa login)."""
    try:
        if fetch_via_next_data(driver):
            return None
    except Exception:
        pass

    visible = visible_text_lower(driver)
    title = (driver.title or "").lower()
    src = page_source_lower(driver)

    if any(m in visible for m in PRIVATE_MARKERS):
        return (
            "Bộ thẻ đang ở CHẾ ĐỘ RIÊNG TƯ — chỉ thành viên lớp Quizlet mới xem được.\n"
            "  → Đăng nhập Quizlet bằng tài khoản đã được duyệt vào lớp (vd. MLN111).\n"
            "  → Mở lại link bộ thẻ; phải thấy flashcard thật, không phải trang 'riêng tư'.\n"
            "  → Chạy lại: python scripts/fetch_quizlet.py <url> --connect 9222"
        )

    if "denied" in title or "px-captcha" in src:
        return (
            "Quizlet chặn bot (PerimeterX). Giải captcha 'Nhấn và giữ' trong Chrome, "
            "hoặc dùng --connect 9222 với Chrome debug."
        )

    if "lỗi ngoài dự tính" in title or "unexpected error" in title:
        return (
            "Quizlet trả lỗi 403. Thử đăng nhập, giải captcha, hoặc dùng --connect 9222."
        )

    return None


def extract_set_id(value: str) -> str:
    value = value.strip()
    if value.isdigit():
        return value
    m = SET_ID_RE.search(value)
    if not m:
        raise ValueError(f"Không tìm thấy set ID trong: {value!r}")
    return m.group(1)


def card_from_item(item: dict[str, Any]) -> dict[str, str]:
    sides = item.get("cardSides") or []
    term = ""
    definition = ""
    if len(sides) >= 1:
        media = sides[0].get("media") or []
        if media:
            term = media[0].get("plainText") or media[0].get("text") or ""
    if len(sides) >= 2:
        media = sides[1].get("media") or []
        if media:
            definition = media[0].get("plainText") or media[0].get("text") or ""
    return {"term": term.strip(), "definition": definition.strip()}


def parse_api_payload(payload: dict[str, Any]) -> tuple[list[dict[str, str]], str | None, int]:
    responses = payload.get("responses") or []
    if not responses:
        return [], None, 0
    resp = responses[0]
    models = resp.get("models") or {}
    raw_items = models.get("studiableItem") or []
    cards = [card_from_item(x) for x in raw_items]
    token = (resp.get("paging") or {}).get("token")
    return cards, token, len(raw_items)


def load_cookies(path: Path) -> dict[str, str]:
    jar = MozillaCookieJar(str(path))
    jar.load(ignore_discard=True, ignore_expires=True)
    return {c.name: c.value for c in jar}


def fetch_via_requests(set_id: str, cookies: dict[str, str]) -> list[dict[str, str]]:
    import requests

    session = requests.Session()
    session.cookies.update(cookies)
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        ),
        "Accept": "application/json",
        "Referer": f"https://quizlet.com/{set_id}/flashcards",
        "Origin": "https://quizlet.com",
    }

    all_cards: list[dict[str, str]] = []
    page = 1
    token: str | None = None

    while True:
        params = {
            "filters[studiableContainerId]": set_id,
            "filters[studiableContainerType]": "1",
            "perPage": str(PER_PAGE),
            "page": str(page),
        }
        if token:
            params["pagingToken"] = token

        r = session.get(API_BASE, params=params, headers=headers, timeout=30)
        if r.status_code != 200:
            body = r.text[:300]
            if "PXVlUfj7uV" in body or "px-captcha" in body:
                raise RuntimeError(
                    "Quizlet chặn request (PerimeterX 403). "
                    "Dùng --connect hoặc --browser thay vì --cookies."
                )
            raise RuntimeError(f"API lỗi HTTP {r.status_code}: {body}")

        batch, token, n = parse_api_payload(r.json())
        all_cards.extend(batch)
        if n < PER_PAGE:
            break
        page += 1

    return [c for c in all_cards if c["term"] or c["definition"]]


def build_api_url(set_id: str, page: int = 1, token: str | None = None) -> str:
    q = (
        f"filters%5BstudiableContainerId%5D={set_id}"
        f"&filters%5BstudiableContainerType%5D=1"
        f"&perPage={PER_PAGE}&page={page}"
    )
    if token:
        q += f"&pagingToken={unquote(token) if '%' in token else token}"
    return f"{API_BASE}?{q}"


def fetch_all_via_browser_js(driver, set_id: str) -> list[dict[str, str]]:
    script = """
    const setId = arguments[0];
    const perPage = arguments[1];
    const cb = arguments[arguments.length - 1];

    (async () => {
      const all = [];
      let page = 1;
      let token = null;
      while (true) {
        let url = `https://quizlet.com/webapi/3.4/studiable-item-documents`
          + `?filters%5BstudiableContainerId%5D=${setId}`
          + `&filters%5BstudiableContainerType%5D=1`
          + `&perPage=${perPage}&page=${page}`;
        if (token) url += `&pagingToken=${encodeURIComponent(token)}`;

        const r = await fetch(url, { credentials: "include" });
        const text = await r.text();
        if (!r.ok) {
          cb(JSON.stringify({ ok: false, status: r.status, body: text.slice(0, 400) }));
          return;
        }
        const data = JSON.parse(text);
        const resp = data.responses?.[0] || {};
        const items = resp.models?.studiableItem || [];
        for (const item of items) {
          const sides = item.cardSides || [];
          const term = sides[0]?.media?.[0]?.plainText || sides[0]?.media?.[0]?.text || "";
          const def = sides[1]?.media?.[0]?.plainText || sides[1]?.media?.[0]?.text || "";
          all.push({ term: String(term).trim(), definition: String(def).trim() });
        }
        if (items.length < perPage) {
          cb(JSON.stringify({ ok: true, cards: all }));
          return;
        }
        token = resp.paging?.token || null;
        page += 1;
        if (!token) {
          cb(JSON.stringify({ ok: true, cards: all }));
          return;
        }
      }
    })();
    """
    raw = driver.execute_async_script(script, set_id, PER_PAGE)
    data = json.loads(raw)
    if not data.get("ok"):
        status = data.get("status")
        body = data.get("body", "")
        raise RuntimeError(f"API trong browser lỗi HTTP {status}: {body[:200]}")
    return [c for c in data["cards"] if c["term"] or c["definition"]]


def fetch_via_next_data(driver) -> list[dict[str, str]]:
    from selenium.webdriver.common.by import By

    els = driver.find_elements(By.ID, "__NEXT_DATA__")
    if not els:
        return []
    payload = json.loads(els[0].get_attribute("textContent") or "{}")
    props = payload.get("props", {}).get("pageProps", {})
    key = props.get("dehydratedReduxStateKey")
    if not key:
        return []
    state = json.loads(key) if isinstance(key, str) else key
    items = (
        state.get("studyModesCommon", {})
        .get("studiableData", {})
        .get("studiableItems", [])
    )
    cards = [card_from_item(x) for x in items]
    return [c for c in cards if c["term"] or c["definition"]]


def page_looks_ready(driver, set_id: str) -> tuple[bool, int, str | None]:
    from selenium.webdriver.common.by import By

    blocked = detect_access_block(driver)
    if blocked:
        raise QuizletAccessError(blocked)

    title = (driver.title or "").lower()
    els = driver.find_elements(By.ID, "__NEXT_DATA__")
    if not els:
        return False, 0, "no_next_data"

    props = json.loads(els[0].get_attribute("textContent") or "{}").get("props", {}).get("pageProps", {})
    err = props.get("errorCode")
    if err and str(err) != "0":
        return False, 0, f"error_{err}"

    cards = fetch_via_next_data(driver)
    if cards:
        return True, len(cards), None

    # title có tên set / flashcards và không phải trang lỗi
    if "quizlet" in title and "lỗi" not in title and "error" not in title:
        return True, 0, None

    return False, 0, "waiting"


def wait_for_quizlet(driver, set_id: str, timeout: int, *, interactive: bool) -> None:
    print(f"Đang chờ trang Quizlet sẵn sàng (tối đa {timeout}s)...")
    if interactive:
        print("Trong Chrome: đăng nhập Quizlet → vào lớp → mở bộ thẻ → giải captcha nếu có.")

    deadline = time.time() + timeout
    last_msg = ""
    while time.time() < deadline:
        blocked = detect_access_block(driver)
        if blocked and interactive:
            print(f"\n⚠ {blocked}\n")
            print("Đang chờ bạn xử lý trong Chrome...")
        elif blocked:
            raise QuizletAccessError(blocked)

        try:
            ready, n, reason = page_looks_ready(driver, set_id)
        except QuizletAccessError:
            if not interactive:
                raise
            ready, n, reason = False, 0, "private_or_blocked"
            time.sleep(2)
            continue

        msg = f"  [{int(deadline - time.time())}s] title={driver.title!r} reason={reason} cards={n}"
        if msg != last_msg:
            print(msg)
            last_msg = msg
        if ready and (n > 0 or reason is None):
            # Chỉ coi sẵn sàng khi đã có thẻ hoặc trang load xong không lỗi
            if n > 0:
                return
            # thử fetch API ngay nếu chưa có next_data
            try:
                probe = fetch_all_via_browser_js(driver, set_id)
                if probe:
                    return
            except RuntimeError:
                pass
        time.sleep(2)

    blocked = detect_access_block(driver)
    if blocked:
        raise QuizletAccessError(blocked)

    raise TimeoutError(
        "Hết thời gian chờ. Kiểm tra:\n"
        "  1. Đã đăng nhập Quizlet và được duyệt vào lớp (bộ riêng tư)?\n"
        "  2. Trang hiển thị flashcard thật, không phải 'Học phần này ở chế độ riêng tư'?\n"
        "  3. Dùng --connect 9222 và tăng --wait 180"
    )


def attach_chrome(port: int):
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options

    opts = Options()
    opts.add_experimental_option("debuggerAddress", f"127.0.0.1:{port}")
    return webdriver.Chrome(options=opts)


def get_chrome_major_version() -> int | None:
    """Đọc phiên bản Chrome cài trên máy (Windows/macOS/Linux)."""
    import shutil
    import subprocess

    candidates = [
        Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe"),
        Path(r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"),
        Path.home() / "AppData/Local/Google/Chrome/Application/chrome.exe",
        Path("/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"),
        Path("/usr/bin/google-chrome"),
        Path("/usr/bin/chromium-browser"),
    ]
    for path in candidates:
        if not path.is_file():
            continue
        if sys.platform == "win32":
            try:
                import winreg

                key = winreg.OpenKey(
                    winreg.HKEY_CURRENT_USER,
                    r"Software\Google\Chrome\BLBeacon",
                )
                version, _ = winreg.QueryValueEx(key, "version")
                winreg.CloseKey(key)
                m = re.search(r"^(\d+)\.", str(version))
                if m:
                    return int(m.group(1))
            except OSError:
                pass
            try:
                version = path.parent / "chrome.VisualElementsManifest.xml"
                if not version.exists():
                    # Last resort: đọc version từ thư mục Application (vd. 149.0.7827.155)
                    for child in path.parent.iterdir():
                        if re.fullmatch(r"\d+\.\d+\.\d+\.\d+", child.name):
                            return int(child.name.split(".")[0])
            except OSError:
                pass
        try:
            out = subprocess.check_output(
                [str(path), "--version"],
                stderr=subprocess.DEVNULL,
                text=True,
                timeout=10,
            )
            m = re.search(r"(\d+)\.", out)
            if m:
                return int(m.group(1))
        except (OSError, subprocess.SubprocessError):
            continue

    chrome = shutil.which("google-chrome") or shutil.which("chromium") or shutil.which("chrome")
    if chrome:
        try:
            out = subprocess.check_output(
                [chrome, "--version"],
                stderr=subprocess.DEVNULL,
                text=True,
                timeout=10,
            )
            m = re.search(r"(\d+)\.", out)
            if m:
                return int(m.group(1))
        except (OSError, subprocess.SubprocessError):
            pass
    return None


def launch_browser(headless: bool, profile_dir: Path | None):
    import undetected_chromedriver as uc

    options = uc.ChromeOptions()
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--window-size=1280,900")
    if profile_dir:
        profile_dir.mkdir(parents=True, exist_ok=True)
        options.add_argument(f"--user-data-dir={profile_dir}")
    if headless:
        options.add_argument("--headless=new")
    version_main = get_chrome_major_version()
    if version_main:
        print(f"Chrome major version: {version_main}")
    return uc.Chrome(options=options, version_main=version_main)


def find_quizlet_tab(driver) -> str | None:
    for handle in driver.window_handles:
        driver.switch_to.window(handle)
        url = driver.current_url or ""
        if "quizlet.com" in url and SET_ID_RE.search(url):
            return url
    return None


def fetch_via_browser(
    set_id: str | None,
    *,
    connect_port: int | None,
    headless: bool,
    wait: int,
    url: str | None,
    profile_dir: Path | None,
) -> tuple[list[dict[str, str]], dict[str, Any]]:
    own_driver = connect_port is None
    driver = attach_chrome(connect_port) if connect_port else launch_browser(headless, profile_dir)
    meta: dict[str, Any] = {"title": "", "url": "", "set_id": set_id}

    try:
        if connect_port:
            tab_url = find_quizlet_tab(driver)
            if tab_url:
                print(f"Đã tìm tab Quizlet: {tab_url}")
                meta["url"] = tab_url
                if not set_id:
                    set_id = extract_set_id(tab_url)
            elif url:
                driver.get(url)
            elif set_id:
                driver.get(f"https://quizlet.com/{set_id}/flashcards")
            else:
                raise ValueError(
                    "Không thấy tab Quizlet. Mở bộ flashcard trong Chrome debug trước, "
                    "hoặc truyền URL/set ID."
                )
        else:
            target = url or f"https://quizlet.com/{set_id}/flashcards"
            print(f"Mở {target}")
            if "quizlet.com" not in target:
                target = f"https://quizlet.com/{set_id}/flashcards"
            # warmup cookie PerimeterX
            driver.get("https://quizlet.com/")
            time.sleep(4)
            driver.get(target)

        if not set_id:
            set_id = extract_set_id(driver.current_url)
        meta["set_id"] = set_id
        meta["url"] = driver.current_url
        meta["title"] = driver.title

        wait_for_quizlet(driver, set_id, wait, interactive=connect_port is not None or not headless)

        try:
            cards = fetch_all_via_browser_js(driver, set_id)
            meta["source"] = "webapi"
        except RuntimeError as e:
            print(f"API thất bại ({e}), thử __NEXT_DATA__...")
            cards = fetch_via_next_data(driver)
            meta["source"] = "next_data"
            if not cards:
                raise

        meta["count"] = len(cards)
        return cards, meta
    finally:
        if own_driver:
            driver.quit()


def _set_run_font(run, *, name: str = "Times New Roman", size=None, bold: bool = False, color=None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = name
    run.font.size = size or Pt(10)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)


def _set_cell_margins(cell, top=40, start=80, bottom=40, end=80) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcMar"))
    if old is not None:
        tc_pr.remove(old)
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def _shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_multiline_text(paragraph, text: str, *, bold: bool = False, size=None, color=None) -> None:
    from docx.enum.text import WD_BREAK

    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        _set_run_font(paragraph.add_run(line), bold=bold, size=size, color=color)


def write_docx(cards: list[dict[str, str]], path: Path, meta: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    FONT = "Times New Roman"
    BODY_SIZE = Pt(10)
    TITLE_SIZE = Pt(14)
    META_SIZE = Pt(9)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE

    title = meta.get("title") or f"Quizlet {meta.get('set_id', '')}"
    title = re.sub(r"\s*[-|]\s*Quizlet\s*$", "", title, flags=re.I).strip()
    title = re.sub(r"^Thẻ ghi nhớ:\s*", "", title, flags=re.I).strip() or "Flashcards Quizlet"

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    _set_run_font(title_p.add_run(title), size=TITLE_SIZE, bold=True)

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(8)
    _set_run_font(
        info_p.add_run(f"{len(cards)} câu · MLN111 · In 2 cột tiết kiệm giấy"),
        size=META_SIZE,
        color=RGBColor(0x55, 0x55, 0x55),
    )

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    col_widths = (Cm(9.2), Cm(7.3))
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    hdr = table.rows[0].cells
    _shade_cell(hdr[0], "D9E2F3")
    _shade_cell(hdr[1], "E2EFDA")

    for idx, label in enumerate(("CÂU HỎI", "ĐÁP ÁN")):
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_margins(hdr[idx], top=60, bottom=60)
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _set_run_font(p.add_run(label), size=Pt(10), bold=True)

    for i, card in enumerate(cards, 1):
        row = table.add_row()
        q_cell, a_cell = row.cells
        for cell in (q_cell, a_cell):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(cell)

        q_para = q_cell.paragraphs[0]
        q_para.paragraph_format.space_before = Pt(0)
        q_para.paragraph_format.space_after = Pt(2)
        q_para.paragraph_format.line_spacing = 1.05
        _set_run_font(q_para.add_run(f"{i}. "), bold=True, size=BODY_SIZE)
        _add_multiline_text(q_para, card["term"], size=BODY_SIZE)

        a_para = a_cell.paragraphs[0]
        a_para.paragraph_format.space_before = Pt(0)
        a_para.paragraph_format.space_after = Pt(2)
        a_para.paragraph_format.line_spacing = 1.05
        _add_multiline_text(
            a_para,
            card["definition"],
            bold=True,
            size=BODY_SIZE,
            color=RGBColor(0x0B, 0x6E, 0x4F),
        )

    doc.save(path)


def parse_cards_from_docx(path: Path) -> tuple[list[dict[str, str]], dict[str, Any]]:
    from docx import Document

    doc = Document(path)
    meta: dict[str, Any] = {"title": doc.paragraphs[0].text.strip() if doc.paragraphs else ""}
    cards: list[dict[str, str]] = []

    if doc.tables:
        table = doc.tables[0]
        for row in table.rows[1:]:
            q_text = row.cells[0].text.strip()
            a_text = row.cells[1].text.strip()
            q_text = re.sub(r"^\d+\.\s*", "", q_text)
            if q_text or a_text:
                cards.append({"term": q_text, "definition": a_text})
        if cards:
            return cards, meta

    current: dict[str, str] = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        q_match = re.match(r"^Câu\s+(\d+)\s*:\s*(.+)$", text, re.I | re.S)
        if q_match:
            if current.get("term") or current.get("definition"):
                cards.append(current)
            current = {"term": q_match.group(2).strip(), "definition": ""}
            continue
        a_match = re.match(r"^Đáp án\s*:\s*(.+)$", text, re.I | re.S)
        if a_match and current:
            current["definition"] = a_match.group(1).strip()
    if current.get("term") or current.get("definition"):
        cards.append(current)
    return cards, meta


def write_output(cards: list[dict[str, str]], path: Path, fmt: str, meta: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fmt == "json":
        payload = {"meta": meta, "cards": cards}
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    elif fmt == "csv":
        with path.open("w", encoding="utf-8-sig", newline="") as f:
            w = csv.DictWriter(f, fieldnames=["term", "definition"])
            w.writeheader()
            w.writerows(cards)
    elif fmt == "txt":
        lines = []
        for i, c in enumerate(cards, 1):
            lines.append(f"{i}. {c['term']}\t{c['definition']}")
        path.write_text("\n".join(lines), encoding="utf-8")
    elif fmt == "docx":
        write_docx(cards, path, meta)
    else:
        raise ValueError(f"Format không hỗ trợ: {fmt}")


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")

    parser = argparse.ArgumentParser(description="Lấy flashcard Quizlet (bypass PerimeterX)")
    parser.add_argument("target", nargs="?", help="URL Quizlet hoặc set ID (vd. 1017891468)")
    parser.add_argument("-o", "--output", type=Path, help="File output")
    parser.add_argument(
        "--format",
        choices=("json", "csv", "txt", "docx"),
        default="docx",
        help="Định dạng output (mặc định: docx)",
    )
    parser.add_argument(
        "--connect",
        type=int,
        metavar="PORT",
        help="Gắn Chrome đang chạy với --remote-debugging-port=PORT (khuyến nghị)",
    )
    parser.add_argument(
        "--browser",
        action="store_true",
        help="Tự mở undetected Chrome (mặc định nếu không --connect/--cookies)",
    )
    parser.add_argument("--cookies", type=Path, help="File cookie Netscape (.txt) export khi ĐÃ đăng nhập Quizlet")
    parser.add_argument(
        "--profile",
        type=Path,
        default=Path.home() / "AppData/Local/ChromeQuizletScrape",
        help="Thư mục profile Chrome (--browser) để giữ đăng nhập",
    )
    parser.add_argument("--headless", action="store_true", help="Chạy Chrome headless (dễ bị chặn)")
    parser.add_argument("--wait", type=int, default=180, help="Giây chờ login/captcha/trang load (mặc định 180)")
    parser.add_argument(
        "--from-json",
        type=Path,
        metavar="FILE",
        help="Chỉ xuất docx/csv/txt từ file JSON đã lấy trước đó (không gọi Quizlet)",
    )
    parser.add_argument(
        "--from-docx",
        type=Path,
        metavar="FILE",
        help="Tạo lại docx (định dạng mới) từ file docx/JSON đã có",
    )
    args = parser.parse_args()

    if args.from_docx:
        if args.from_docx.suffix.lower() == ".json":
            payload = json.loads(args.from_docx.read_text(encoding="utf-8-sig"))
            cards = payload.get("cards") or payload
            meta = payload.get("meta") or {"source": "from_json"}
        else:
            cards, meta = parse_cards_from_docx(args.from_docx)
        if not cards:
            print("Không đọc được câu hỏi từ file nguồn.", file=sys.stderr)
            return 1
        output = args.output or args.from_docx
        if args.format != "docx":
            output = args.output or (ROOT / "data" / f"quizlet_export.{args.format}")
        else:
            output = args.output or args.from_docx
        write_output(cards, output, args.format if args.format != "docx" else "docx", meta)
        print(f"\n✓ Đã xuất {len(cards)} thẻ → {output}")
        return 0

    if args.from_json:
        payload = json.loads(args.from_json.read_text(encoding="utf-8-sig"))
        cards = payload.get("cards") or payload
        meta = payload.get("meta") or {"source": "from_json"}
        output = args.output or (ROOT / "data" / f"quizlet_export.{args.format}")
        write_output(cards, output, args.format, meta)
        print(f"\n✓ Đã xuất {len(cards)} thẻ → {output}")
        return 0

    set_id: str | None = None
    url: str | None = None
    if args.target:
        if args.target.isdigit():
            set_id = args.target
        elif "quizlet.com" in args.target:
            url = args.target
            set_id = extract_set_id(args.target)
        else:
            set_id = extract_set_id(args.target)

    default_name = f"quizlet_{set_id or 'export'}.{args.format}"
    output = args.output or (ROOT / "data" / default_name)

    meta: dict[str, Any] = {"set_id": set_id}

    try:
        if args.cookies:
            if not set_id:
                raise ValueError("Cần set ID hoặc URL khi dùng --cookies")
            print(f"Lấy qua API + cookies ({args.cookies})...")
            cards = fetch_via_requests(set_id, load_cookies(args.cookies))
            meta["source"] = "cookies_api"
            meta["count"] = len(cards)
        else:
            if not args.connect and not args.browser and not args.target:
                parser.print_help()
                print("\nBộ MLN111 là RIÊNG TƯ — cần đăng nhập + vào lớp trước:")
                print('  1. chrome.exe --remote-debugging-port=9222 --user-data-dir="%LOCALAPPDATA%\\ChromeDebug"')
                print("  2. Đăng nhập Quizlet → xin vào lớp MLN111")
                print("  3. Mở link bộ thẻ, thấy flashcard thật")
                print('  4. python scripts/fetch_quizlet.py "URL" --connect 9222 -o data/quizlet_mln111.json')
                return 1

            use_connect = args.connect is not None
            if not use_connect and not args.browser:
                args.browser = True

            cards, meta = fetch_via_browser(
                set_id,
                connect_port=args.connect,
                headless=args.headless,
                wait=args.wait,
                url=url,
                profile_dir=args.profile,
            )

        write_output(cards, output, args.format, meta)
        print(f"\n✓ Đã lấy {len(cards)} thẻ → {output}")
        if cards:
            print(f"  Mẫu: {cards[0]['term'][:60]} → {cards[0]['definition'][:60]}")
        return 0

    except Exception as e:
        print(f"\n✗ Lỗi: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
