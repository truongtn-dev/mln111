#!/usr/bin/env python3
"""Debug docx structure for parser issues."""
import re
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
OPTION_START = re.compile(r"^([a-eA-E])[\.\)]\s*(.*)$")
QUESTION_START = re.compile(
    r"^(?:Câu\s*(?:mở đầu|[0-9]+(?:\.[0-9]+)?)|C[0-9]+(?:\.[0-9]+)?)\s*[:：\.]\s*(.+)",
    re.I,
)


def para_bold_ratio(para):
    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return 0.0
    bold = sum(len(r.text) for r in runs if r.bold)
    total = sum(len(r.text) for r in runs)
    return bold / total if total else 0.0


def show_para(para, idx):
    t = para.text.strip()
    if not t:
        return
    br = para_bold_ratio(para)
    opt = OPTION_START.match(t)
    q = QUESTION_START.match(t)
    tag = []
    if q:
        tag.append("Q")
    if opt:
        tag.append(f"OPT-{opt.group(1)}")
    if br >= 0.35:
        tag.append(f"BOLD={br:.0%}")
    runs = [(r.text[:40], r.bold) for r in para.runs if r.text.strip()]
    print(f"[{idx}] {'/'.join(tag) or 'TXT'}: {t[:100]}")
    if runs and len(runs) > 1:
        print(f"      runs: {runs[:6]}")


def main():
    for path in sorted(ROOT.glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        print(f"\n{'='*60}\n{path.name}\n{'='*60}")
        doc = Document(path)
        for i, para in enumerate(doc.paragraphs):
            t = para.text.strip()
            if not t:
                continue
            if re.search(r"Câu\s*25|C25:|C1:|Sản xuất xã hội|hình thức vận động|hình thức", t, re.I):
                for j in range(max(0, i - 1), min(len(doc.paragraphs), i + 10)):
                    show_para(doc.paragraphs[j], j)
                print("---")


if __name__ == "__main__":
    main()
